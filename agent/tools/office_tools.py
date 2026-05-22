# ============================================================
# agent/tools/office_tools.py
# Tools untuk membuat dokumen Word dan Excel — JARVIS V5
# Mendukung markdown parsing untuk .docx
# ============================================================

import os
import re
from typing import Any

from agent.config import resolve_path

# Import python-docx
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Import openpyxl
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False


def _resolve_save_path(filename: str, save_path: str = "") -> str:
    """
    Tentukan path penyimpanan file.

    Args:
        filename:  Nama file (bisa termasuk path)
        save_path: Folder tujuan (opsional)

    Returns:
        Path absolut lengkap untuk menyimpan file
    """
    if save_path:
        folder = resolve_path(save_path)
        os.makedirs(folder, exist_ok=True)
        return os.path.join(folder, os.path.basename(filename))
    elif os.path.isabs(filename):
        os.makedirs(os.path.dirname(filename), exist_ok=True)
        return filename
    else:
        base = resolve_path("")
        return os.path.join(base, filename)


def _add_inline_bold(paragraph: Any, text: str) -> None:
    """
    Parse teks dengan **bold** dan tambahkan run ke paragraph docx.

    Args:
        paragraph: Objek paragraph docx
        text:      Teks dengan kemungkinan **bold** markers
    """
    # Pecah berdasarkan **...**
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:  # Bagian di dalam **...**
            run.bold = True


def create_docx(filename: str, content: str, save_path: str = "") -> str:
    """
    Buat dokumen Word (.docx) dari teks dengan markdown formatting.

    Markdown yang didukung:
    - # Heading 1
    - ## Heading 2
    - ### Heading 3
    - **teks bold**
    - - bullet list
    - 1. numbered list
    - | col | col | → tabel (baris 1 = header bold)
    - --- → page break

    Args:
        filename:  Nama file output (akan ditambah .docx jika belum ada)
        content:   Konten dengan markdown formatting
        save_path: Folder tujuan penyimpanan (opsional)

    Returns:
        Pesan sukses dengan path file atau pesan error
    """
    if not HAS_DOCX:
        return "Error: python-docx tidak terinstall. Jalankan: pip install python-docx"

    try:
        # Pastikan ekstensi .docx
        if not filename.endswith(".docx"):
            filename += ".docx"

        output_path = _resolve_save_path(filename, save_path)

        doc = Document()

        # ── Atur margin A4 standar ────────────────────────────
        from docx.shared import Cm
        section = doc.sections[0]
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(2.54)
        section.right_margin  = Cm(2.54)
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)

        lines = content.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i]

            # ── Page break ───────────────────────────────────
            if line.strip() == "---":
                doc.add_page_break()
                i += 1
                continue

            # ── Heading 1 ────────────────────────────────────
            if line.startswith("# ") and not line.startswith("## "):
                h = doc.add_heading(line[2:].strip(), level=1)
                h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
                i += 1
                continue

            # ── Heading 2 ────────────────────────────────────
            if line.startswith("## ") and not line.startswith("### "):
                h = doc.add_heading(line[3:].strip(), level=2)
                i += 1
                continue

            # ── Heading 3 ────────────────────────────────────
            if line.startswith("### "):
                h = doc.add_heading(line[4:].strip(), level=3)
                i += 1
                continue

            # ── Tabel markdown ───────────────────────────────
            if line.strip().startswith("|") and "|" in line:
                # Kumpulkan semua baris tabel
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    row_line = lines[i].strip()
                    # Skip baris separator (|---|---|)
                    if not re.match(r'^[\|\s\-:]+$', row_line):
                        table_lines.append(row_line)
                    i += 1

                if table_lines:
                    # Parse kolom
                    rows_data = []
                    for tl in table_lines:
                        cols = [c.strip() for c in tl.strip("|").split("|")]
                        rows_data.append(cols)

                    max_cols = max(len(r) for r in rows_data)
                    table = doc.add_table(rows=len(rows_data), cols=max_cols)
                    table.style = "Table Grid"

                    for row_idx, row_data in enumerate(rows_data):
                        row = table.rows[row_idx]
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < max_cols:
                                cell = row.cells[col_idx]
                                cell.text = cell_text
                                # Baris header: bold + background biru muda
                                if row_idx == 0:
                                    for run in cell.paragraphs[0].runs:
                                        run.bold = True
                                    # Set background
                                    tc = cell._tc
                                    tcPr = tc.get_or_add_tcPr()
                                    shd = OxmlElement("w:shd")
                                    shd.set(qn("w:fill"), "BDD7EE")
                                    shd.set(qn("w:color"), "auto")
                                    shd.set(qn("w:val"), "clear")
                                    tcPr.append(shd)
                continue

            # ── Bullet list ──────────────────────────────────
            if line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(style="List Bullet")
                _add_inline_bold(p, line[2:].strip())
                i += 1
                continue

            # ── Numbered list ────────────────────────────────
            if re.match(r'^\d+\. ', line):
                p = doc.add_paragraph(style="List Number")
                text = re.sub(r'^\d+\. ', '', line).strip()
                _add_inline_bold(p, text)
                i += 1
                continue

            # ── Baris kosong ─────────────────────────────────
            if not line.strip():
                doc.add_paragraph()
                i += 1
                continue

            # ── Paragraf biasa dengan inline bold ────────────
            p = doc.add_paragraph()
            _add_inline_bold(p, line.strip())
            i += 1

        doc.save(output_path)
        size = os.path.getsize(output_path)
        return f"✅ Dokumen Word dibuat: {output_path} ({size/1024:.1f} KB)"

    except Exception as e:
        return f"Error create_docx: {e}"


def create_xlsx(
    filename: str,
    rows: list,
    save_path: str = "",
    sheet_name: str = "Sheet1"
) -> str:
    """
    Buat spreadsheet Excel (.xlsx) dari data list of lists.

    Fitur:
    - Baris pertama otomatis bold + background biru
    - Auto column width berdasarkan konten
    - Border tipis di semua cell

    Args:
        filename:   Nama file output (akan ditambah .xlsx jika belum ada)
        rows:       List of lists, baris pertama = header
                    Contoh: [["Nama","Usia"], ["Budi","25"], ["Sari","30"]]
        save_path:  Folder tujuan penyimpanan (opsional)
        sheet_name: Nama sheet (default: Sheet1)

    Returns:
        Pesan sukses dengan path file atau pesan error
    """
    if not HAS_XLSX:
        return "Error: openpyxl tidak terinstall. Jalankan: pip install openpyxl"

    try:
        # Pastikan ekstensi .xlsx
        if not filename.endswith(".xlsx"):
            filename += ".xlsx"

        output_path = _resolve_save_path(filename, save_path)

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name

        # ── Style header ─────────────────────────────────────
        header_font    = Font(bold=True, color="FFFFFF", size=11)
        header_fill    = PatternFill("solid", fgColor="2E75B6")
        header_align   = Alignment(horizontal="center", vertical="center", wrap_text=True)

        # ── Style data ───────────────────────────────────────
        data_font      = Font(size=10)
        data_align     = Alignment(vertical="center", wrap_text=True)
        alt_fill       = PatternFill("solid", fgColor="DEEAF1")  # Baris genap

        # ── Border ───────────────────────────────────────────
        thin_border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )

        # ── Isi data ─────────────────────────────────────────
        col_widths: dict[int, int] = {}

        for row_idx, row_data in enumerate(rows, start=1):
            for col_idx, value in enumerate(row_data, start=1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.border = thin_border

                if row_idx == 1:
                    # Baris header
                    cell.font   = header_font
                    cell.fill   = header_fill
                    cell.alignment = header_align
                else:
                    # Baris data dengan alternating color
                    cell.font = data_font
                    cell.alignment = data_align
                    if row_idx % 2 == 0:
                        cell.fill = alt_fill

                # Hitung lebar kolom maksimal
                content_len = len(str(value)) if value is not None else 0
                if col_widths.get(col_idx, 0) < content_len:
                    col_widths[col_idx] = content_len

        # ── Set lebar kolom (auto width) ─────────────────────
        for col_idx, width in col_widths.items():
            col_letter = get_column_letter(col_idx)
            # Min 8, max 50, tambah padding 2
            ws.column_dimensions[col_letter].width = max(8, min(50, width + 2))

        # ── Freeze baris header ──────────────────────────────
        ws.freeze_panes = "A2"

        wb.save(output_path)
        size = os.path.getsize(output_path)
        total_rows = len(rows) - 1 if rows else 0
        return (
            f"✅ Spreadsheet Excel dibuat: {output_path} "
            f"({size/1024:.1f} KB, {total_rows} baris data)"
        )

    except Exception as e:
        return f"Error create_xlsx: {e}"
