"""
╔══════════════════════════════════════════════════════════════╗
║        JARVIS V44 OMNI — GLOBAL DEEPSCAN ARCHITECTURE        ║
║   Parallel Intelligence Gathering + ReAct Agent Reasoning    ║
╚══════════════════════════════════════════════════════════════╝
"""

import webbrowser
import subprocess
import re
import json
import os
import sys
import time
import shutil
import itertools
import concurrent.futures
import urllib.parse
import urllib.request
from datetime import datetime
from typing import Any, Dict, List, Optional

# ─── DEPENDENCIES ──────────────────────────────────────────────────────────────
from groq import Groq
from colorama import init, Fore, Style

try:
    import pyfiglet
    HAS_FIGLET = True
except ImportError:
    HAS_FIGLET = False

try:
    from tavily import TavilyClient
    HAS_TAVILY = True
except ImportError:
    HAS_TAVILY = False

try:
    import shodan
    HAS_SHODAN = True
except ImportError:
    HAS_SHODAN = False

try:
    import psutil
    HAS_PSUTIL = True
except ImportError:
    HAS_PSUTIL = False

try:
    from docx import Document as DocxDocument
    from openpyxl import Workbook
    HAS_OFFICE = True
except ImportError:
    HAS_OFFICE = False

init(autoreset=True)

# ─── API KEYS ──────────────────────────────────────────────────────────────────
# Isi API key kamu di sini atau gunakan environment variable
GROQ_API_KEY   = os.environ.get("GROQ_API_KEY", "ISI_GROQ_API_KEY_KAMU")
TAVILY_API_KEY = os.environ.get("TAVILY_API_KEY", "ISI_TAVILY_API_KEY_KAMU")
SHODAN_API_KEY = os.environ.get("SHODAN_API_KEY", "ISI_SHODAN_API_KEY_KAMU")

MODEL       = "llama-3.3-70b-versatile"
MAX_STEPS   = 8
WORK_DIR    = os.path.dirname(os.path.abspath(__file__))

# ─── CLIENT INIT ───────────────────────────────────────────────────────────────
try:
    groq_client = Groq(api_key=GROQ_API_KEY)
except Exception as e:
    print(f"{Fore.RED}[!] Groq init error: {e}")
    sys.exit(1)

tavily_client = TavilyClient(api_key=TAVILY_API_KEY) if HAS_TAVILY else None
shodan_api    = shodan.Shodan(SHODAN_API_KEY) if HAS_SHODAN else None

# ══════════════════════════════════════════════════════════════════════════════
#  LAYER 1: CORE UTILITIES
# ══════════════════════════════════════════════════════════════════════════════

def parse_json_robust(text: str) -> Optional[Dict]:
    if not text:
        return None
    try:
        return json.loads(text)
    except Exception:
        pass

    m = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', text, re.DOTALL)
    if m:
        try:
            return json.loads(m.group(1))
        except Exception:
            pass

    m = re.search(r'\{.*\}', text, re.DOTALL)
    if m:
        try:
            return json.loads(m.group(0))
        except Exception:
            pass

    return None


def call_llm(messages: List[Dict], temperature: float = 0.15) -> str:
    try:
        resp = groq_client.chat.completions.create(
            model=MODEL,
            messages=messages,
            temperature=temperature,
            max_tokens=4096,
        )
        return resp.choices[0].message.content or ""
    except Exception as e:
        return f"[LLM ERROR] {e}"


# ══════════════════════════════════════════════════════════════════════════════
#  LAYER 2: TOOLS
# ══════════════════════════════════════════════════════════════════════════════

def tool_web_search(query: str, max_results: int = 5) -> Dict:
    if not tavily_client:
        return {"error": "Tavily not available. Install: pip install tavily-python"}
    try:
        r = tavily_client.search(query=query, max_results=max_results, include_raw_content=False)
        return {
            "results": [
                {"title": x.get("title", ""), "url": x.get("url", ""), "content": x.get("content", "")}
                for x in r.get("results", [])
            ]
        }
    except Exception as e:
        return {"error": str(e)}


def tool_system_info() -> Dict:
    if not HAS_PSUTIL:
        return {"error": "psutil not available. Install: pip install psutil"}
    try:
        cpu_freq = psutil.cpu_freq()
        mem  = psutil.virtual_memory()
        disk = psutil.disk_usage(os.path.splitdrive(WORK_DIR)[0] + "\\")
        net  = psutil.net_io_counters()
        bat  = psutil.sensors_battery()
        return {
            "cpu_percent":    psutil.cpu_percent(interval=0.5),
            "cpu_cores":      psutil.cpu_count(logical=True),
            "cpu_freq_mhz":   round(cpu_freq.current, 1) if cpu_freq else None,
            "ram_total_gb":   round(mem.total / 1024**3, 2),
            "ram_used_gb":    round(mem.used  / 1024**3, 2),
            "ram_percent":    mem.percent,
            "disk_total_gb":  round(disk.total / 1024**3, 2),
            "disk_used_gb":   round(disk.used  / 1024**3, 2),
            "disk_percent":   disk.percent,
            "net_sent_mb":    round(net.bytes_sent / 1024**2, 2),
            "net_recv_mb":    round(net.bytes_recv / 1024**2, 2),
            "battery_pct":    round(bat.percent, 1) if bat else None,
            "plugged_in":     bat.power_plugged if bat else None,
            "boot_time":      datetime.fromtimestamp(psutil.boot_time()).strftime("%Y-%m-%d %H:%M:%S"),
            "platform":       sys.platform,
        }
    except Exception as e:
        return {"error": str(e)}


def tool_list_processes(top_n: int = 10) -> List[Dict]:
    if not HAS_PSUTIL:
        return [{"error": "psutil not available"}]
    try:
        procs = []
        for p in psutil.process_iter(["pid", "name", "cpu_percent", "memory_percent"]):
            try:
                procs.append(p.info)
            except Exception:
                pass
        procs.sort(key=lambda x: x.get("cpu_percent") or 0, reverse=True)
        return procs[:top_n]
    except Exception as e:
        return [{"error": str(e)}]


def tool_shodan_scan(target: str) -> Dict:
    if not shodan_api:
        return {"error": "Shodan not available. Install: pip install shodan"}
    try:
        host = shodan_api.host(target)
        return {
            "ip":          host.get("ip_str"),
            "org":         host.get("org"),
            "country":     host.get("country_name"),
            "os":          host.get("os"),
            "ports":       host.get("ports", []),
            "vulns":       list(host.get("vulns", {}).keys()),
            "hostnames":   host.get("hostnames", []),
            "last_update": host.get("last_update"),
        }
    except Exception as e:
        return {"error": str(e)}


def tool_open_browser(url: str) -> Dict:
    try:
        webbrowser.open(url)
        return {"status": "opened", "url": url}
    except Exception as e:
        return {"error": str(e)}


def tool_run_command(cmd: str, timeout: int = 30) -> Dict:
    try:
        result = subprocess.run(
            cmd, shell=True, capture_output=True, text=True,
            timeout=timeout, encoding="utf-8", errors="replace"
        )
        return {
            "stdout":     result.stdout[:3000],
            "stderr":     result.stderr[:500],
            "returncode": result.returncode,
        }
    except subprocess.TimeoutExpired:
        return {"error": f"Timeout after {timeout}s"}
    except Exception as e:
        return {"error": str(e)}


def tool_save_text(filename: str, content: str) -> Dict:
    try:
        path = os.path.join(WORK_DIR, filename)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        return {"status": "saved", "path": path}
    except Exception as e:
        return {"error": str(e)}


def tool_read_file(filename: str) -> Dict:
    try:
        path = os.path.join(WORK_DIR, filename)
        if not os.path.exists(path):
            return {"error": f"File not found: {filename}"}
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read(8000)
        return {"content": content, "path": path}
    except Exception as e:
        return {"error": str(e)}


def tool_create_docx(filename: str, content: str) -> Dict:
    if not HAS_OFFICE:
        return {"error": "python-docx not installed. Run: pip install python-docx"}
    try:
        if not filename.endswith(".docx"):
            filename += ".docx"
        path = os.path.join(WORK_DIR, filename)
        doc = DocxDocument()
        for line in content.split("\n"):
            if line.startswith("### "):
                doc.add_heading(line[4:], 3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], 2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], 1)
            else:
                doc.add_paragraph(line)
        doc.save(path)
        return {"status": "created", "path": path}
    except Exception as e:
        return {"error": str(e)}


def tool_create_xlsx(filename: str, rows: List[List]) -> Dict:
    if not HAS_OFFICE:
        return {"error": "openpyxl not installed. Run: pip install openpyxl"}
    try:
        if not filename.endswith(".xlsx"):
            filename += ".xlsx"
        path = os.path.join(WORK_DIR, filename)
        wb = Workbook()
        ws = wb.active
        for row in rows:
            ws.append(row)
        wb.save(path)
        return {"status": "created", "path": path}
    except Exception as e:
        return {"error": str(e)}


# ══════════════════════════════════════════════════════════════════════════════
#  LAYER 3: PARALLEL INTEL GATHERING + HTML REPORT
# ══════════════════════════════════════════════════════════════════════════════

def gather_intel_parallel(target: str) -> Dict:
    queries = [
        f"{target} profil biodata lengkap",
        f"{target} pekerjaan jabatan perusahaan",
        f"{target} alamat lokasi domisili",
        f"{target} akun sosial media instagram twitter",
        f"{target} berita kontroversi terbaru 2025",
        f"{target} email kontak informasi",
    ]

    results = {}

    def search_one(q):
        return q, tool_web_search(q, max_results=3)

    with concurrent.futures.ThreadPoolExecutor(max_workers=6) as ex:
        futures = {ex.submit(search_one, q): q for q in queries}
        for future in concurrent.futures.as_completed(futures):
            try:
                q, res = future.result()
                results[q] = res
            except Exception as e:
                results[futures[future]] = {"error": str(e)}

    return results


def synthesize_intel(target: str, raw_data: Dict) -> Dict:
    data_str = json.dumps(raw_data, ensure_ascii=False)[:6000]

    prompt = f"""Kamu adalah analis intelijen profesional. Dari data mentah berikut tentang "{target}",
ekstrak dan sintesis informasi terstruktur dalam format JSON.

Data mentah:
{data_str}

Kembalikan HANYA JSON dengan struktur berikut (null jika tidak diketahui):
{{
  "nama": "nama lengkap",
  "lokasi": "kota, negara",
  "pekerjaan": "jabatan",
  "perusahaan": "nama perusahaan/organisasi",
  "email": null,
  "sosmed": {{"instagram": null, "twitter": null, "linkedin": null, "facebook": null}},
  "koneksi": ["nama orang/entitas terkait"],
  "risk_indicators": ["indikator risiko jika ada"],
  "aktivitas_terbaru": "deskripsi aktivitas terbaru",
  "tentang": "ringkasan singkat 2-3 kalimat",
  "confidence": "HIGH atau MEDIUM atau LOW",
  "sources": ["url1", "url2"]
}}

Hanya JSON, tanpa teks tambahan."""

    resp = call_llm([{"role": "user", "content": prompt}], temperature=0.1)
    parsed = parse_json_robust(resp)

    if not parsed:
        return {
            "nama": target, "lokasi": None, "pekerjaan": None,
            "perusahaan": None, "email": None,
            "sosmed": {}, "koneksi": [], "risk_indicators": [],
            "aktivitas_terbaru": "Synthesis failed",
            "tentang": f"Raw data collected from {len(raw_data)} search layers.",
            "confidence": "LOW", "sources": [],
        }

    return parsed


def generate_html_report(target: str, intel: Dict) -> str:
    ts       = datetime.now().strftime("%Y-%m-%d %H:%M")
    conf     = intel.get("confidence", "LOW")
    conf_clr = {"HIGH": "#00ff41", "MEDIUM": "#ff8800", "LOW": "#ff0040"}.get(conf, "#7d8590")

    # vis.js nodes / edges
    nodes = [{"id": 1, "label": target, "shape": "dot", "size": 35,
               "color": {"background": "#00ff41", "border": "#00ff41"}}]
    edges = []
    nid = 2
    for person in (intel.get("koneksi") or [])[:8]:
        nodes.append({"id": nid, "label": person, "shape": "dot", "size": 20,
                       "color": {"background": "#0088ff", "border": "#0088ff"}})
        edges.append({"from": 1, "to": nid})
        nid += 1

    # HTML blocks
    def info_row(icon, key, val):
        return f'<div class="info-row"><span class="info-key">{icon} {key}</span><span class="info-val">{val or "N/A"}</span></div>'

    sosmed_html = "".join(
        info_row("🔗", p.title(), h)
        for p, h in (intel.get("sosmed") or {}).items() if h
    )

    email_row = info_row("📧", "Email", intel.get("email")) if intel.get("email") else ""

    risk_html = "".join(
        f'<div style="color:#ff4444;font-size:12px;margin:4px 0;">⚠ {r}</div>'
        for r in (intel.get("risk_indicators") or [])
    ) or '<div style="color:#7d8590;font-size:12px;">No risk indicators found.</div>'

    sources_html = "".join(
        f'<a href="{u}" target="_blank" style="display:block;color:#0088ff;font-size:11px;'
        f'margin:2px 0;overflow:hidden;white-space:nowrap;text-overflow:ellipsis;">'
        f'{urllib.parse.urlparse(u).netloc}</a>'
        for u in (intel.get("sources") or [])[:8]
    )

    nodes_json = json.dumps(nodes)
    edges_json = json.dumps(edges)
    target_upper = target.upper()
    nama    = intel.get("nama") or target
    lokasi  = intel.get("lokasi")
    job     = intel.get("pekerjaan")
    company = intel.get("perusahaan")
    aktiv   = intel.get("aktivitas_terbaru") or "N/A"
    tentang = intel.get("tentang") or ""

    return f"""<!DOCTYPE html>
<html lang="id">
<head>
  <meta charset="UTF-8">
  <title>JARVIS INTEL | {target}</title>
  <script src="https://unpkg.com/vis-network/standalone/umd/vis-network.min.js"></script>
  <style>
    *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{ background: #010409; color: #c9d1d9; font-family: 'Consolas', 'Courier New', monospace; }}
    .topbar {{ background: #0d1117; border-bottom: 1px solid #21262d; padding: 12px 20px; display: flex; justify-content: space-between; align-items: center; }}
    .topbar-title {{ color: #00ff41; font-size: 13px; letter-spacing: 2px; }}
    .grid {{ display: grid; grid-template-columns: 280px 1fr 280px; gap: 1px; height: calc(100vh - 45px); background: #21262d; }}
    .panel {{ background: #010409; padding: 16px; overflow-y: auto; }}
    .panel-title {{ font-size: 11px; letter-spacing: 3px; color: #7d8590; border-bottom: 1px solid #21262d; padding-bottom: 8px; margin-bottom: 12px; text-transform: uppercase; }}
    .target-name {{ color: #00ff41; font-size: 18px; font-weight: bold; margin-bottom: 4px; }}
    .badge {{ display: inline-block; padding: 2px 8px; border-radius: 3px; font-size: 10px; margin: 2px; }}
    .info-row {{ display: flex; gap: 8px; margin: 6px 0; font-size: 12px; }}
    .info-key {{ color: #7d8590; min-width: 80px; }}
    .info-val {{ color: #e6edf3; }}
    .about-text {{ font-size: 12px; line-height: 1.7; color: #8b949e; margin-top: 10px; }}
    #graph {{ width: 100%; height: 100%; min-height: 400px; }}
    .graph-wrap {{ height: calc(100% - 30px); }}
    a {{ color: #0088ff; text-decoration: none; }}
    a:hover {{ text-decoration: underline; }}
    ::-webkit-scrollbar {{ width: 4px; }}
    ::-webkit-scrollbar-track {{ background: #010409; }}
    ::-webkit-scrollbar-thumb {{ background: #30363d; }}
  </style>
</head>
<body>
  <div class="topbar">
    <div class="topbar-title">⬡ JARVIS V44 // INTEL REPORT // {target_upper}</div>
    <div style="font-size:11px;">
      <span style="color:{conf_clr}">CONFIDENCE: {conf}</span>
      &nbsp;|&nbsp;
      <span style="color:#7d8590">{ts}</span>
    </div>
  </div>
  <div class="grid">
    <div class="panel">
      <div class="panel-title">Intel Dossier</div>
      <div class="target-name">{nama}</div>
      <div style="margin:10px 0;">
        <span class="badge" style="background:#0d1117;border:1px solid {conf_clr};color:{conf_clr};">{conf} CONF</span>
      </div>
      {info_row("📍", "Lokasi", lokasi)}
      {info_row("💼", "Job", job)}
      {info_row("🏢", "Company", company)}
      {email_row}
      {sosmed_html}
      <div style="margin-top:12px;padding-top:12px;border-top:1px solid #21262d;">
        <div class="panel-title" style="margin-bottom:8px;">Aktivitas Terbaru</div>
        <div class="about-text">{aktiv}</div>
      </div>
      <div class="about-text" style="margin-top:12px;">{tentang}</div>
    </div>
    <div class="panel" style="padding:10px;">
      <div class="panel-title">Relation Radar</div>
      <div class="graph-wrap"><div id="graph"></div></div>
    </div>
    <div class="panel">
      <div class="panel-title" style="color:#ff0040;">Risk Indicators</div>
      {risk_html}
      <div style="margin-top:16px;">
        <div class="panel-title">Data Sources</div>
        {sources_html}
      </div>
      <div style="margin-top:16px;">
        <div class="panel-title">Sosial Media</div>
        {sosmed_html or '<div style="color:#7d8590;font-size:12px;">Tidak ditemukan.</div>'}
      </div>
    </div>
  </div>
  <script>
    var nodes = new vis.DataSet({nodes_json});
    var edges = new vis.DataSet({edges_json});
    var options = {{
      nodes: {{ font: {{ color: '#00ff41', size: 11 }}, borderWidth: 1, shadow: true }},
      edges: {{ width: 1, smooth: {{ type: 'continuous' }} }},
      physics: {{ barnesHut: {{ gravitationalConstant: -4000, springLength: 120 }} }}
    }};
    new vis.Network(document.getElementById('graph'), {{nodes, edges}}, options);
  </script>
</body>
</html>"""


# ══════════════════════════════════════════════════════════════════════════════
#  LAYER 4: REACT AGENT
# ══════════════════════════════════════════════════════════════════════════════

TOOL_DEFS = {
    "web_search": {
        "desc": "Cari informasi di internet via Tavily",
        "params": {"query": "string", "max_results": "int (opsional, default 5)"},
    },
    "system_info": {
        "desc": "Dapatkan info sistem laptop (CPU, RAM, disk, baterai, network)",
        "params": {},
    },
    "list_processes": {
        "desc": "Tampilkan proses yang berjalan berurutan berdasarkan CPU",
        "params": {"top_n": "int (opsional, default 10)"},
    },
    "shodan_scan": {
        "desc": "Scan IP address via Shodan untuk port dan kerentanan",
        "params": {"target": "string IP address"},
    },
    "open_browser": {
        "desc": "Buka URL di browser",
        "params": {"url": "string"},
    },
    "run_command": {
        "desc": "Jalankan perintah shell Windows dan ambil outputnya",
        "params": {"cmd": "string", "timeout": "int (opsional, default 30)"},
    },
    "save_text": {
        "desc": "Simpan teks ke file",
        "params": {"filename": "string", "content": "string"},
    },
    "read_file": {
        "desc": "Baca isi file dari direktori kerja",
        "params": {"filename": "string"},
    },
    "create_docx": {
        "desc": "Buat dokumen Word (.docx). Gunakan # ## ### untuk heading",
        "params": {"filename": "string", "content": "string"},
    },
    "create_xlsx": {
        "desc": "Buat spreadsheet Excel (.xlsx) dari list of lists",
        "params": {"filename": "string", "rows": "list of lists"},
    },
    "intel_scan": {
        "desc": "Scan intelijen mendalam (paralel) pada orang/perusahaan, buat laporan HTML",
        "params": {"target": "string nama target"},
    },
    "final_answer": {
        "desc": "Berikan jawaban final ke user",
        "params": {"answer": "string"},
    },
}

SYSTEM_PROMPT = f"""Kamu adalah JARVIS V44 OMNI, asisten AI canggih berbasis ReAct (Reason + Act).

Pada setiap langkah, kembalikan HANYA JSON dengan format:
{{
  "thought": "analisis situasi dan rencanamu",
  "action": "nama_tool",
  "action_input": {{parameter}}
}}

Saat sudah punya jawaban final:
{{
  "thought": "selesai",
  "action": "final_answer",
  "action_input": {{"answer": "jawaban lengkap dalam Bahasa Indonesia"}}
}}

Tools tersedia:
{json.dumps(TOOL_DEFS, indent=2, ensure_ascii=False)}

Aturan:
- Gunakan tool paling relevan untuk setiap tugas
- Info laptop/sistem → system_info
- Pencarian info → web_search
- Intel mendalam → intel_scan
- Jangan ulangi tool yang sama dengan input sama
- Jawab Bahasa Indonesia kecuali diminta lain
- Hanya kembalikan JSON, tidak ada teks lain di luar JSON
"""


def execute_tool(action: str, action_input: Dict) -> str:
    try:
        if action == "web_search":
            res = tool_web_search(
                action_input.get("query", ""),
                int(action_input.get("max_results", 5))
            )
        elif action == "system_info":
            res = tool_system_info()
        elif action == "list_processes":
            res = tool_list_processes(int(action_input.get("top_n", 10)))
        elif action == "shodan_scan":
            res = tool_shodan_scan(action_input.get("target", ""))
        elif action == "open_browser":
            res = tool_open_browser(action_input.get("url", ""))
        elif action == "run_command":
            res = tool_run_command(
                action_input.get("cmd", ""),
                int(action_input.get("timeout", 30))
            )
        elif action == "save_text":
            res = tool_save_text(
                action_input.get("filename", "output.txt"),
                action_input.get("content", "")
            )
        elif action == "read_file":
            res = tool_read_file(action_input.get("filename", ""))
        elif action == "create_docx":
            res = tool_create_docx(
                action_input.get("filename", "document"),
                action_input.get("content", "")
            )
        elif action == "create_xlsx":
            res = tool_create_xlsx(
                action_input.get("filename", "spreadsheet"),
                action_input.get("rows", [])
            )
        elif action == "intel_scan":
            target = action_input.get("target", "")
            print(f"\n{Fore.YELLOW}  [*] Parallel intel scan: {target}{Style.RESET_ALL}")
            raw   = gather_intel_parallel(target)
            print(f"{Fore.YELLOW}  [*] Synthesizing with LLM...{Style.RESET_ALL}")
            intel = synthesize_intel(target, raw)
            html  = generate_html_report(target, intel)
            ts    = datetime.now().strftime("%H%M%S")
            safe  = re.sub(r"[^\w]", "_", target)[:20]
            fname = f"INTEL_{safe}_{ts}.html"
            fpath = os.path.join(WORK_DIR, fname)
            with open(fpath, "w", encoding="utf-8") as f:
                f.write(html)
            webbrowser.open("file:///" + fpath.replace("\\", "/"))
            res = {"status": "Intel report generated & opened in browser", "file": fname, "summary": intel}
        else:
            res = {"error": f"Unknown tool: {action}"}

        return json.dumps(res, ensure_ascii=False, default=str)[:4000]

    except Exception as e:
        return json.dumps({"error": str(e)})


def react_agent(user_input: str) -> str:
    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "user",   "content": user_input},
    ]

    for step in range(MAX_STEPS):
        raw    = call_llm(messages)
        parsed = parse_json_robust(raw)

        if not parsed:
            return raw  # LLM gave plain text — return as-is

        thought      = parsed.get("thought", "")
        action       = parsed.get("action", "")
        action_input = parsed.get("action_input", {})

        if thought:
            print(f"  {Fore.BLUE}[THINK]{Style.RESET_ALL} {thought[:120]}")

        if action == "final_answer":
            return action_input.get("answer", str(action_input))

        if not action:
            return raw

        inp_preview = json.dumps(action_input, ensure_ascii=False)[:80]
        print(f"  {Fore.CYAN}[ACT]  {Style.RESET_ALL}{action}({inp_preview})")

        observation  = execute_tool(action, action_input)
        obs_preview  = observation[:150] + "..." if len(observation) > 150 else observation
        print(f"  {Fore.GREEN}[OBS]  {Style.RESET_ALL}{obs_preview}")

        messages.append({"role": "assistant", "content": raw})
        messages.append({"role": "user",      "content": f"Observation: {observation}"})

    # Force final answer after max steps
    messages.append({"role": "user", "content": "Berikan jawaban final sekarang berdasarkan semua observasi."})
    return call_llm(messages)


# ══════════════════════════════════════════════════════════════════════════════
#  LAYER 5: MAIN LOOP
# ══════════════════════════════════════════════════════════════════════════════

def print_banner():
    if HAS_FIGLET:
        print(Fore.GREEN + pyfiglet.figlet_format("JARVIS  V44", font="slant"))
    else:
        print(Fore.GREEN + r"""
     ___ ___    ____  ____ __  __ ____  _____  _   __ _  _    _
    |_ _|   \  / _  ||  _ \\ \/ /|_ _|/ ___/ | | / // || \  | |
     | || |) || |_| || |_) |\  /  | | \___  \ | |/ / | |  \_| |
    |___|___/  \___/ |____/ |__|  |_| /____/  |___/  |_|\______|
""")
    print(f"{Fore.CYAN}  OMNI — Global DeepScan + ReAct Intelligence Architecture")
    print(f"{Fore.YELLOW}  LLM  : {MODEL}")
    print(f"{Fore.YELLOW}  Tools: {len(TOOL_DEFS) - 1} active  |  WorkDir: {WORK_DIR}")
    print(f"{Fore.MAGENTA}  Ketik 'help' untuk daftar perintah | 'exit' untuk keluar")
    print(f"{Fore.WHITE}{'═' * 65}{Style.RESET_ALL}\n")


HELP_TEXT = f"""
{Fore.CYAN}━━━ JARVIS V44 OMNI — BANTUAN ━━━{Style.RESET_ALL}

{Fore.GREEN}Percakapan Bebas:{Style.RESET_ALL}
  Ketik pertanyaan apa saja dalam Bahasa Indonesia / Inggris

{Fore.GREEN}Perintah Cepat:{Style.RESET_ALL}
  sysinfo            Info sistem laptop (CPU, RAM, disk, baterai)
  proses             Daftar proses berjalan teratas
  intel <target>     Scan intelijen mendalam pada orang/perusahaan
  scan <ip>          Scan IP via Shodan
  cmd <perintah>     Jalankan perintah shell Windows
  help               Tampilkan bantuan ini
  exit / quit        Keluar

{Fore.GREEN}Contoh:{Style.RESET_ALL}
  > sysinfo
  > intel Elon Musk
  > scan 8.8.8.8
  > cmd ipconfig /all
  > cek berapa RAM yang terpakai
  > cariin berita AI terbaru 2025
  > buatkan file excel daftar belanja bulanan
  > buat dokumen Word laporan keuangan sederhana
"""


def format_response(text: str) -> str:
    lines = []
    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("# "):
            lines.append(Fore.CYAN + Style.BRIGHT + line + Style.RESET_ALL)
        elif stripped.startswith("## ") or stripped.startswith("### "):
            lines.append(Fore.CYAN + line + Style.RESET_ALL)
        elif stripped.startswith(("- ", "* ", "• ")):
            lines.append(Fore.WHITE + line + Style.RESET_ALL)
        elif any(w in stripped.lower() for w in ("error", "gagal", "tidak bisa", "failed")):
            lines.append(Fore.RED + line + Style.RESET_ALL)
        elif stripped.startswith(("✓", "✔", "berhasil", "sukses")):
            lines.append(Fore.GREEN + line + Style.RESET_ALL)
        else:
            lines.append(line)
    return "\n".join(lines)


def main():
    print_banner()

    # Dependency status
    checks = [
        ("Tavily",  HAS_TAVILY,  Fore.GREEN, Fore.RED),
        ("Shodan",  HAS_SHODAN,  Fore.GREEN, Fore.YELLOW),
        ("psutil",  HAS_PSUTIL,  Fore.GREEN, Fore.YELLOW),
        ("Office",  HAS_OFFICE,  Fore.GREEN, Fore.YELLOW),
        ("figlet",  HAS_FIGLET,  Fore.GREEN, Fore.YELLOW),
    ]
    parts = []
    for name, ok, c_ok, c_no in checks:
        sym = "✓" if ok else "✗"
        clr = c_ok if ok else c_no
        parts.append(f"{clr}{name} {sym}{Style.RESET_ALL}")
    print("  Status: " + "  ".join(parts) + "\n")

    while True:
        try:
            user_input = input(f"{Fore.GREEN}┌─[JARVIS]─❯{Style.RESET_ALL} ").strip()
        except (KeyboardInterrupt, EOFError):
            print(f"\n{Fore.YELLOW}[JARVIS] Sampai jumpa!{Style.RESET_ALL}")
            break

        if not user_input:
            continue

        low = user_input.lower()

        if low in ("exit", "quit", "keluar", "bye"):
            print(f"{Fore.YELLOW}[JARVIS] Sampai jumpa!{Style.RESET_ALL}")
            break

        if low == "help":
            print(HELP_TEXT)
            continue

        # Shortcut expansions
        if low == "sysinfo":
            user_input = "Tampilkan info sistem laptop saya secara lengkap dan rapi dalam format yang mudah dibaca"
        elif low == "proses":
            user_input = "Tampilkan 10 proses yang paling banyak menggunakan CPU saat ini"
        elif low.startswith("intel "):
            user_input = f"Lakukan intel scan mendalam pada target: {user_input[6:].strip()}"
        elif low.startswith("scan "):
            user_input = f"Scan IP {user_input[5:].strip()} menggunakan Shodan dan tampilkan hasilnya"
        elif low.startswith("cmd "):
            user_input = f"Jalankan perintah shell ini dan tampilkan outputnya: {user_input[4:].strip()}"

        print(f"\n{Fore.MAGENTA}{'─' * 55}{Style.RESET_ALL}")
        t0 = time.time()

        try:
            response = react_agent(user_input)
            elapsed  = time.time() - t0

            print(f"\n{Fore.GREEN}┌─[JARVIS RESPONSE]{'─' * 38}{Style.RESET_ALL}")
            print(format_response(response))
            print(f"{Fore.GREEN}└{'─' * 56}{Style.RESET_ALL}")
            print(f"  {Fore.WHITE}⏱ {elapsed:.1f}s{Style.RESET_ALL}\n")

        except KeyboardInterrupt:
            print(f"\n{Fore.YELLOW}[!] Dihentikan.{Style.RESET_ALL}")
        except Exception as e:
            print(f"\n{Fore.RED}[!] Error: {e}{Style.RESET_ALL}")


if __name__ == "__main__":
    main()
